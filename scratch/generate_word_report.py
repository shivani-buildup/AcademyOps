"""
AcademyOps — Final Project Report Generator (Word .docx)
Generates a highly detailed, professional final project report using mostly bullet points and max 3 tables.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# ─── Paths ──────────────────────────────────────────────────────────────────
BASE_DIR = r"c:\Users\Admin\Desktop\AcademyOps"
OUTPUT_PATH = os.path.join(BASE_DIR, "AcademyOps_Final_Project_Report_v5.docx")

IMG_FUNNEL_CHART = os.path.join(BASE_DIR, "funnel_chart.png")
IMG_SOURCE_CHART = os.path.join(BASE_DIR, "source_comparison.png")

# ─── Color Palette ──────────────────────────────────────────────────────────
COLOR_PRIMARY = RGBColor(79, 70, 229)
COLOR_SECONDARY = RGBColor(59, 130, 246)
COLOR_DARK = RGBColor(15, 23, 42)
COLOR_MUTED = RGBColor(71, 85, 105)

# ─── Helper Functions ───────────────────────────────────────────────────────
def add_simple_table(doc, headers, rows, col_widths=None):
    """Add a simple, standard table without custom XML to ensure it opens safely in MS Word."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(header)
        p = cell.paragraphs[0]
        p.runs[0].bold = True
        p.runs[0].font.name = "Calibri"
        p.runs[0].font.size = Pt(11)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(cell_text)
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(11)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph("")  # Spacer
    return table

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = COLOR_PRIMARY if level <= 2 else COLOR_DARK
        run.font.name = 'Calibri'
    return heading

def add_body_text(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        p.add_run(text).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(50, 50, 50)
    p.paragraph_format.space_after = Pt(10)
    return p

def add_image_safe(doc, img_path, width=Inches(5.5), caption=None):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = COLOR_MUTED
    else:
        add_body_text(doc, f"[Image not found: {os.path.basename(img_path)}]", italic=True, color=COLOR_MUTED)

# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Default Font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = COLOR_DARK

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── TITLE PAGE ───
for _ in range(6): doc.add_paragraph("")

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("AcademyOps\n")
run.font.size = Pt(44)
run.font.bold = True
run.font.color.rgb = COLOR_PRIMARY
run.font.name = "Calibri"

run2 = title_p.add_run("Lead-to-Enrollment Management System\n")
run2.font.size = Pt(20)
run2.font.color.rgb = COLOR_SECONDARY
run2.font.name = "Calibri"

title_p.add_run("\nProfessional Final Project Report\nFor EasySkill Career Academy (ECA)").font.size = Pt(14)
doc.add_paragraph("")
div_p = doc.add_paragraph()
div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
div_p.add_run("━" * 50).font.color.rgb = COLOR_PRIMARY
doc.add_paragraph("")
doc.add_paragraph("")

# TABLE 1: Project Details
details = [
    ("Project Name", "AcademyOps - Lead Management System"),
    ("Developer / Student", "Shivani Kosambiya"),
    ("Target Organization", "EasySkill Career Academy (ECA)"),
    ("Core Technologies", "Python, FastAPI, SQLAlchemy, Streamlit, Plotly"),
    ("Report Date", "June 2026"),
    ("Project Version", "1.0.0 Production Build"),
]
add_simple_table(doc, ["Project Attribute", "Description"], details, col_widths=[2.5, 4.0])

doc.add_page_break()

# ─── TOC ───
add_heading_styled(doc, "Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Problem Statement & Objectives",
    "3. System Architecture & Flow",
    "4. Technology Stack & Justification",
    "5. Database Design & Models",
    "6. Data Ingestion & Normalization (WP-02)",
    "7. RESTful API Implementation (WP-07)",
    "8. Intent Classification Engine (WP-08)",
    "9. Operations Dashboard & Analytics (WP-05/06)",
    "10. Automated Testing & Quality Assurance",
    "11. CI/CD Pipeline & Deployment Strategy",
    "12. UI/UX Design Screenshots",
    "13. Conclusion & Future Enhancements"
]
for item in toc_items:
    add_body_text(doc, item)
doc.add_page_break()


# ─── CHAPTER 1 ───
add_heading_styled(doc, "1. Executive Summary", level=1)
add_body_text(doc, 
    "AcademyOps is a robust, full-stack, data-driven Lead-to-Enrollment Management System meticulously engineered for the EasySkill Career Academy (ECA). "
    "In the modern educational landscape, managing prospective student inquiries across multiple digital channels (Facebook, Instagram, Google, Organic Website traffic, etc.) is a complex challenge. "
    "AcademyOps addresses this by providing a unified, centralized platform that captures, validates, tracks, and analyzes leads through a formally defined 6-stage sales pipeline."
)
add_body_text(doc, 
    "This system serves as a bridge between marketing efforts and counselor operations. It eliminates manual data entry errors through automated CSV ingestion pipelines, standardizes data formats, and utilizes a custom RESTful API to manage the entire lifecycle of a student inquiry."
)
add_body_text(doc, 
    "Furthermore, the project integrates an intelligent Intent Classifier leveraging Natural Language Processing (NLP) rules to automatically categorize student messages and suggest appropriate responses, significantly reducing the cognitive load on counselors. "
    "Coupled with a real-time Streamlit analytics dashboard, the academy's management gains immediate visibility into funnel performance, conversion metrics, and marketing source effectiveness."
)


# ─── CHAPTER 2 ───
add_heading_styled(doc, "2. Problem Statement & Objectives", level=1)
add_heading_styled(doc, "2.1 Problem Statement", level=2)
add_body_text(doc, "Prior to the implementation of AcademyOps, EasySkill Career Academy faced several critical operational bottlenecks:")
add_bullet(doc, "Leads were manually recorded in disconnected spreadsheets, leading to duplicate entries and formatting inconsistencies.", bold_prefix="Data Fragmentation: ")
add_bullet(doc, "A lack of real-time pipeline tracking meant management could not determine how many leads were actively being nurtured.", bold_prefix="Zero Pipeline Visibility: ")
add_bullet(doc, "Counselors wasted hours manually categorizing routine inquiries (e.g., questions about fees, timings) and typing repetitive responses.", bold_prefix="Inefficient Communication: ")
add_bullet(doc, "Without centralized data, generating accurate reports on which marketing channels yielded the highest enrollment rates was impossible.", bold_prefix="Lack of Analytics: ")

add_heading_styled(doc, "2.2 Core Objectives", level=2)
add_body_text(doc, "To resolve these challenges, the project was designed with the following strategic objectives:")
add_bullet(doc, "Develop an automated data ingestion and sanitization pipeline to clean messy marketing CSV data.", bold_prefix="Objective 1: ")
add_bullet(doc, "Architect a scalable relational database with proper schema definitions and indexes for optimal querying.", bold_prefix="Objective 2: ")
add_bullet(doc, "Expose a secure, well-documented REST API using FastAPI to handle all backend CRUD operations.", bold_prefix="Objective 3: ")
add_bullet(doc, "Implement a rule-based Intent Classifier to automate the categorization of incoming student text messages.", bold_prefix="Objective 4: ")
add_bullet(doc, "Deploy a real-time, interactive operations dashboard for advanced statistical and funnel analysis.", bold_prefix="Objective 5: ")
add_bullet(doc, "Ensure enterprise-grade reliability through rigorous automated testing and continuous integration (CI/CD).", bold_prefix="Objective 6: ")


# ─── CHAPTER 3 ───
add_heading_styled(doc, "3. System Architecture & Flow", level=1)
add_body_text(doc, "The AcademyOps platform follows a highly modular, multi-tier Layered Architecture. This ensures a clear separation of concerns, making the system maintainable, testable, and scalable.")

add_heading_styled(doc, "3.1 System Architecture Flowchart", level=2)
p_arch = doc.add_paragraph("⬇️ [ અહી SYSTEM ARCHITECTURE FLOWCHART ની ઈમેજ ADD કરો ] ⬇️")
p_arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_arch.runs[0].font.color.rgb = RGBColor(220, 38, 38) # Red color for visibility
p_arch.runs[0].bold = True
doc.add_paragraph("")

add_heading_styled(doc, "3.2 Architectural Layers", level=2)
add_bullet(doc, "Built using Streamlit, this layer consumes data and presents interactive analytics (Plotly charts, KPI metrics) to the end-user. The FastAPI Swagger interface also serves as a developer-facing presentation layer.", bold_prefix="Presentation Layer (Frontend): ")
add_bullet(doc, "Constructed with FastAPI, this layer exposes secure RESTful endpoints. It validates incoming JSON payloads using Pydantic schemas before passing them deeper into the system.", bold_prefix="API Layer (Routing): ")
add_bullet(doc, "This is the brain of the application. It contains the logic for the NLP Intent Classifier, the Data Importer/Cleanser, and the core rules governing pipeline stage transitions.", bold_prefix="Business Logic Layer (Services): ")
add_bullet(doc, "Utilizing SQLAlchemy ORM, this layer abstracts raw SQL queries into Pythonic object operations. It utilizes a Repository Pattern to handle data persistence securely.", bold_prefix="Data Access Layer (ORM): ")
add_bullet(doc, "The foundational tier utilizing SQLite (with an easy migration path to PostgreSQL). It handles raw data storage with strict constraints and indexing.", bold_prefix="Database Layer: ")

add_heading_styled(doc, "3.3 Data Flow Lifecycle", level=2)
add_body_text(doc, "1. Raw CSV data is ingested via the Importer Script, which sanitizes text and normalizes phone numbers.")
add_body_text(doc, "2. Valid data is committed to the database; invalid data is pushed to a Quarantine file.")
add_body_text(doc, "3. The Dashboard queries the API/Database in real-time to render visual funnel metrics.")


# ─── CHAPTER 4 ───
add_heading_styled(doc, "4. Technology Stack & Justification", level=1)
add_body_text(doc, "The technology stack was carefully selected to prioritize performance, developer velocity, and maintainability. Below is the detailed breakdown of the technologies used:")

add_bullet(doc, "Chosen for its simplicity, massive ecosystem, and excellent support for data analysis and machine learning tasks.", bold_prefix="Python (Core Language): ")
add_bullet(doc, "Selected over Flask/Django for the backend API because of its asynchronous capabilities, exceptional speed, and out-of-the-box automatic Swagger/OpenAPI documentation.", bold_prefix="FastAPI (Web Framework): ")
add_bullet(doc, "An ASGI web server that provides the necessary async performance to serve the FastAPI application concurrently.", bold_prefix="Uvicorn (ASGI Server): ")
add_bullet(doc, "The industry-standard ORM for Python. It prevents SQL injection attacks and allows seamless switching between SQLite for development and PostgreSQL for production.", bold_prefix="SQLAlchemy (Database ORM): ")
add_bullet(doc, "Provides robust data validation and settings management using Python type annotations. It ensures that the API only accepts properly formatted data.", bold_prefix="Pydantic (Validation): ")
add_bullet(doc, "A rapid prototyping framework that allowed the creation of a beautiful, interactive frontend dashboard purely in Python without writing HTML/JS.", bold_prefix="Streamlit (Dashboard UI): ")
add_bullet(doc, "Used for rendering highly interactive and visually appealing graphs (like the Funnel Chart and Source Donut Chart) in the dashboard.", bold_prefix="Plotly & Pandas (Analytics): ")
add_bullet(doc, "Chosen for automated unit and integration testing. httpx is used alongside it to simulate asynchronous API requests during tests.", bold_prefix="Pytest & httpx (Testing): ")


# ─── CHAPTER 5 ───
add_heading_styled(doc, "5. Database Design & Models", level=1)
add_body_text(doc, "The database design is centralized around the core 'leads' entity. Stringent constraints and indexes are applied at the database level to ensure absolute data integrity.")

# TABLE 2: Database Schema (Table 2 of 3)
add_heading_styled(doc, "5.1 Leads Table Schema", level=2)
schema_rows = [
    ["id", "INTEGER", "PRIMARY KEY, AUTO INCREMENT", "Unique internal identifier for the lead."],
    ["name", "TEXT", "NOT NULL", "The sanitized full name of the prospective student."],
    ["phone", "TEXT", "NOT NULL, UNIQUE", "The validated phone number. Indexed for fast lookup and deduplication."],
    ["source", "TEXT", "DEFAULT 'Unknown'", "The normalized marketing channel (e.g., Facebook, Instagram, Google)."],
    ["stage", "TEXT", "NOT NULL, INDEXED", "Current position in the 6-stage sales pipeline."],
    ["notes", "TEXT", "NULLABLE", "Optional unstructured text for counselor observations."],
    ["created_at", "DATETIME", "DEFAULT NOW()", "Timestamp indicating when the lead entered the system."],
    ["updated_at", "DATETIME", "DEFAULT NOW()", "Timestamp automatically updated upon any record modification."]
]
add_simple_table(doc, ["Column", "Data Type", "Constraints", "Description"], schema_rows, col_widths=[1.0, 1.2, 2.0, 2.5])


# ─── CHAPTER 6 ───
add_heading_styled(doc, "6. Data Ingestion & Normalization (WP-02)", level=1)
add_body_text(doc, "Marketing data rarely arrives clean. The Data Ingestion Pipeline (WP-02) acts as the system's gatekeeper, ensuring that only pristine data enters the database.")

add_heading_styled(doc, "6.1 Data Ingestion Pipeline Flowchart", level=2)
p_pipe = doc.add_paragraph("⬇️ [ અહી DATA INGESTION PIPELINE FLOWCHART ની ઈમેજ ADD કરો ] ⬇️")
p_pipe.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_pipe.runs[0].font.color.rgb = RGBColor(220, 38, 38)
p_pipe.runs[0].bold = True
doc.add_paragraph("")

add_heading_styled(doc, "6.2 Normalization Rules", level=2)
add_bullet(doc, "Trims leading/trailing whitespace and converts the string to Title Case (e.g., '  john DOE ' becomes 'John Doe').", bold_prefix="Name Formatting: ")
add_bullet(doc, "Uses a predefined mapping dictionary to standardize variations. For example, 'fb', 'FB', and 'facebook' are all mapped to 'Facebook'. Unrecognized sources are capitalized.", bold_prefix="Source Mapping: ")
add_bullet(doc, "Regular expressions (Regex) are used to strip invalid characters. Phone numbers containing alphabetic characters are strictly rejected.", bold_prefix="Phone Validation: ")

add_heading_styled(doc, "6.3 The Quarantine System", level=2)
add_body_text(doc, "The system employs a strict 'fail-safe' mechanism. If a lead fails validation (e.g., missing a name or having a completely invalid phone number) or is a duplicate, it is NOT discarded. Instead, it is written to a specialized `quarantine.csv` file along with a `rejection_reason` column. This allows administrators to manually review and fix broken leads without halting the main ingestion process.")


# ─── CHAPTER 7 ───
add_heading_styled(doc, "7. RESTful API Implementation (WP-07)", level=1)
add_body_text(doc, "The backend is powered by a comprehensive REST API built with FastAPI. It handles all Create, Read, Update, and Delete (CRUD) operations and enforces business rules.")

add_heading_styled(doc, "7.1 Pipeline Stages Enforced by the API", level=2)
add_body_text(doc, "The API strictly validates that a lead's stage must belong to one of the following enumerated values:")
add_bullet(doc, "Lead has been imported but no action has been taken.", bold_prefix="1. New: ")
add_bullet(doc, "A counselor has initiated communication.", bold_prefix="2. Contacted: ")
add_bullet(doc, "The student's requirements match the academy's offerings.", bold_prefix="3. Qualified: ")
add_bullet(doc, "The student has attended a trial class or demo session.", bold_prefix="4. Demo: ")
add_bullet(doc, "The student has successfully paid fees and joined.", bold_prefix="5. Enrolled: ")
add_bullet(doc, "The lead is no longer interested or unreachable.", bold_prefix="6. Lost: ")

add_heading_styled(doc, "7.2 Core Endpoints Overview", level=2)
add_bullet(doc, "Retrieves a paginated list of leads. Supports query parameters to filter by 'source' or 'stage'.", bold_prefix="GET /api/v1/leads: ")
add_bullet(doc, "Accepts a JSON payload (validated via Pydantic) to create a new lead entry.", bold_prefix="POST /api/v1/leads: ")
add_bullet(doc, "Allows counselors to move a lead through the funnel (e.g., from 'New' to 'Contacted').", bold_prefix="PATCH /api/v1/leads/{id}/stage: ")
add_bullet(doc, "Provides real-time aggregations (total leads, conversion rates) directly to the dashboard.", bold_prefix="GET /api/v1/stats: ")


# ─── CHAPTER 8 ───
add_heading_styled(doc, "8. Intent Classification Engine (WP-08)", level=1)
add_body_text(doc, "To streamline communication, AcademyOps features an automated NLP Intent Classifier. When a student sends a message, the system analyzes the text to determine the core intent and suggests a response.")

add_heading_styled(doc, "8.1 Classifier Logic Flowchart", level=2)
p_class = doc.add_paragraph("⬇️ [ અહી INTENT CLASSIFIER FLOWCHART ની ઈમેજ ADD કરો ] ⬇️")
p_class.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_class.runs[0].font.color.rgb = RGBColor(220, 38, 38)
p_class.runs[0].bold = True
doc.add_paragraph("")

add_heading_styled(doc, "8.2 Classification Categories", level=2)
add_bullet(doc, "Triggered by keywords like 'cost', 'fee', 'price'. Suggests a pre-written template explaining course fees and recommends moving the lead to 'Contacted'.", bold_prefix="Intent: Fees — ")
add_bullet(doc, "Triggered by 'duration', 'months', 'schedule'. Suggests a template detailing batch timings.", bold_prefix="Intent: Timing — ")
add_bullet(doc, "Triggered by 'degree', 'background', 'requirement'. Suggests a template explaining prerequisites and recommends moving the lead to 'Qualified'.", bold_prefix="Intent: Eligibility — ")
add_bullet(doc, "Triggered by 'stop', 'unsubscribe', 'no thanks'. Acknowledges the request and recommends moving the lead immediately to the 'Lost' stage.", bold_prefix="Intent: Not Interested — ")
add_bullet(doc, "A catch-all for messages that do not match specific rules. Suggests a generic polite response stating a counselor will call them shortly.", bold_prefix="Intent: Other (Fallback) — ")


# ─── CHAPTER 9 ───
add_heading_styled(doc, "9. Operations Dashboard & Analytics (WP-05/06)", level=1)
add_body_text(doc, "The Streamlit-powered dashboard serves as the central command center for academy managers. It translates raw database rows into actionable visual intelligence.")

add_heading_styled(doc, "9.1 Key Dashboard Features", level=2)
add_bullet(doc, "Displays top-level metrics instantly: Total Leads, Overall Enrollment Rate (%), Active Pipeline count, and Drop-off Rate.", bold_prefix="KPI Cards: ")
add_bullet(doc, "An interactive Plotly funnel chart visually depicting the drop-off at each stage of the sales pipeline (New → Contacted → Qualified → Demo → Enrolled).", bold_prefix="Pipeline Funnel Visualization: ")
add_bullet(doc, "A dynamic donut chart breaking down the volume of leads originating from Facebook, Instagram, Google, etc., allowing marketing budgets to be optimized.", bold_prefix="Source Distribution: ")
add_bullet(doc, "The backend analytics engine includes a 2-Proportion Z-Test implementation. This allows administrators to statistically prove if one marketing channel (e.g., Facebook) is genuinely performing better than another (e.g., Instagram), rather than relying on guesswork.", bold_prefix="Statistical Significance Testing: ")


# ─── CHAPTER 10 ───
add_heading_styled(doc, "10. Automated Testing & Quality Assurance", level=1)
add_body_text(doc, "Enterprise software requires rigorous testing. AcademyOps utilizes the Pytest framework to ensure high reliability across all components.")

add_heading_styled(doc, "10.1 Testing Strategy", level=2)
add_bullet(doc, "Tests individual functions in isolation. For example, verifying that the regex in the phone validation function correctly rejects alphabetic characters.", bold_prefix="Unit Testing: ")
add_bullet(doc, "Tests the interaction between layers. The test suite spins up an in-memory SQLite database, uses the FastAPI TestClient to send mock HTTP requests, and verifies that the database updates correctly without affecting production data.", bold_prefix="Integration Testing: ")
add_bullet(doc, "The test suite covers edge cases such as attempting to insert duplicate phone numbers, requesting non-existent lead IDs, and providing invalid pipeline stages.", bold_prefix="Edge Case Coverage: ")


# ─── CHAPTER 11 ───
add_heading_styled(doc, "11. CI/CD Pipeline & Deployment Strategy", level=1)
add_body_text(doc, "To maintain code quality in a collaborative environment, a Continuous Integration (CI) pipeline was implemented using GitHub Actions.")

# TABLE 3: Work Package Summary (Table 3 of 3)
add_heading_styled(doc, "11.1 Work Package Implementation Summary", level=2)
add_body_text(doc, "The project was executed in a structured, incremental manner using Work Packages (WP):")

wp_rows = [
    ["WP-01", "Project Setup & Virtual Environment", "Completed"],
    ["WP-02", "Data Ingestion Pipeline & Normalization", "Completed"],
    ["WP-03", "Relational Database Design", "Completed"],
    ["WP-04", "Automated Testing Suite (Pytest)", "Completed"],
    ["WP-05", "Analytics & Statistical Engine", "Completed"],
    ["WP-06", "Streamlit Operations Dashboard", "Completed"],
    ["WP-07", "FastAPI REST API Backend", "Completed"],
    ["WP-08", "NLP Intent Classifier", "Completed"],
    ["WP-09", "GitHub Actions CI/CD Pipeline", "Completed"]
]
add_simple_table(doc, ["Work Package", "Feature Description", "Status"], wp_rows, col_widths=[1.5, 3.5, 1.0])


# ─── CHAPTER 12 ───
doc.add_page_break()
add_heading_styled(doc, "12. UI/UX Design Screenshots", level=1)
add_body_text(doc, "This section is dedicated to the visual presentation of the AcademyOps system. The following placeholders are reserved for the high-fidelity screenshots of the graphical user interfaces, demonstrating the user experience designed for academy counselors and administrators.")
add_body_text(doc, "(Please insert the relevant screenshots below the respective headings.)", italic=True, color=COLOR_MUTED)

add_heading_styled(doc, "12.1 Operations Dashboard - KPI & Funnel View", level=2)
p1 = doc.add_paragraph("⬇️ [ અહી MAIN DASHBOARD ની ઈમેજ ADD કરો ] ⬇️")
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.runs[0].font.color.rgb = RGBColor(220, 38, 38)
p1.runs[0].bold = True
doc.add_paragraph("")
doc.add_paragraph("")

add_heading_styled(doc, "12.2 Analytics - Source Distribution Chart", level=2)
p2 = doc.add_paragraph("⬇️ [ અહી SOURCE DISTRIBUTION CHART ની ઈમેજ ADD કરો ] ⬇️")
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.runs[0].font.color.rgb = RGBColor(220, 38, 38)
p2.runs[0].bold = True
doc.add_paragraph("")
doc.add_paragraph("")

add_heading_styled(doc, "12.3 FastAPI Interactive Swagger Documentation", level=2)
p3 = doc.add_paragraph("⬇️ [ અહી SWAGGER UI ની ઈમેજ ADD કરો ] ⬇️")
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.runs[0].font.color.rgb = RGBColor(220, 38, 38)
p3.runs[0].bold = True
doc.add_paragraph("")
doc.add_paragraph("")


# ─── CHAPTER 13 ───
doc.add_page_break()
add_heading_styled(doc, "13. Conclusion & Future Enhancements", level=1)

add_heading_styled(doc, "13.1 Conclusion", level=2)
add_body_text(doc, "AcademyOps successfully achieves all defined objectives, delivering a highly professional, production-ready system for EasySkill Career Academy. By automating data ingestion, enforcing strict data integrity rules, providing a high-performance REST API, and surfacing actionable insights through an interactive dashboard, the system dramatically increases operational efficiency.")
add_body_text(doc, "The integration of NLP for message classification represents a significant leap forward, transforming the platform from a simple database GUI into an intelligent assistant that actively helps counselors close enrollments faster.")

add_heading_styled(doc, "13.2 Future Enhancements", level=2)
add_body_text(doc, "While the current system is robust, future iterations could include:")
add_bullet(doc, "Implementing Machine Learning models (e.g., Scikit-Learn) for predictive lead scoring to identify high-probability enrollments.", bold_prefix="Predictive Lead Scoring: ")
add_bullet(doc, "Connecting the API directly to WhatsApp Business or Twilio for fully automated, real-time student interaction.", bold_prefix="WhatsApp API Integration: ")
add_bullet(doc, "Adding OAuth2 / JWT based authentication to secure endpoints and provide role-based access control (RBAC) for different counselors.", bold_prefix="Security & Authentication: ")

# Final footer
doc.add_paragraph("")
doc.add_paragraph("")
div_p2 = doc.add_paragraph()
div_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
div_p2.add_run("━" * 50).font.color.rgb = COLOR_PRIMARY

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("END OF REPORT").font.bold = True
footer.runs[0].font.color.rgb = COLOR_MUTED

doc.save(OUTPUT_PATH)
print("Professional Detailed Report generated successfully!")
